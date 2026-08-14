# -*- coding: utf-8 -*-
"""生成 ceyuan skill 五问题诊断报告 docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 颜色
RED = RGBColor(0xC0, 0x00, 0x00)      # 结论强调红
DARK = RGBColor(0x22, 0x22, 0x22)     # 正文深灰黑
GRAY = RGBColor(0x66, 0x66, 0x66)     # 次级灰
BLUE = RGBColor(0x1F, 0x4E, 0x79)     # 标题蓝
GREEN = RGBColor(0x1E, 0x7A, 0x46)     # 成立/通过绿

doc = Document()

# 页面边距
for s in doc.sections:
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.4); s.right_margin = Cm(2.4)

def set_font(run, name_zh='微软雅黑', name_en='Microsoft YaHei', size=10.5, color=DARK, bold=False):
    run.font.name = name_en
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    rf.set(qn('w:ascii'), name_en); rf.set(qn('w:hAnsi'), name_en)
    rf.set(qn('w:eastAsia'), name_zh)

def para(text, size=10.5, color=DARK, bold=False, align=None, space_after=6, space_before=0, indent=0):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.35
    if indent: p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    return p

def rich_para(segments, space_after=6, space_before=0, indent=0):
    """segments: list of (text, size, color, bold)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.35
    if indent: p.paragraph_format.left_indent = Cm(indent)
    for (t, s, c, b) in segments:
        run = p.add_run(t); set_font(run, size=s, color=c, bold=b)
    return p

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text); set_font(run, size=16, color=BLUE, bold=True)
    # 底部加线
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), '1F4E79')
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text); set_font(run, size=13, color=BLUE, bold=True)
    return p

def bullet(text, size=10.5, color=DARK, bold=False, marker='· '):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(marker + text); set_font(run, size=size, color=color, bold=bold)
    return p

def verdict(label, text):
    """结论判定行：红标签 + 结论"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.35
    r1 = p.add_run(label + '  '); set_font(r1, size=11, color=RED, bold=True)
    r2 = p.add_run(text); set_font(r2, size=11, color=DARK, bold=False)
    return p

def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htxt)
        set_font(run, size=10, color=RGBColor(0xFF,0xFF,0xFF), bold=True)
        # 表头底色
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'1F4E79')
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            set_font(run, size=9.5, color=DARK)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

# ============ 封面标题 ============
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run = p.add_run('策元 ceyuan skill 五问题诊断报告'); set_font(run, size=20, color=BLUE, bold=True)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('以第一性原理逐条判断「是否需要修改」及「设计者自身不足」')
set_font(run, size=11, color=GRAY)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(14)
run = p.add_run('2026-08-14 · 内部诊断 · 结论先行版')
set_font(run, size=9.5, color=GRAY)

# ============ 结论先行 ============
h1('〇、结论先行（判词）')

rich_para([('五个问题全部成立，全部需要修改。', 11.5, RED, True),
           ('但不都是同一种病——前四个是「方法论结构性缺陷」（P0），第五个是「交付格式缺陷」（P1）。', 11, DARK, False)])

para('一句话概括五个问题的共同根因：', size=11, bold=True, space_after=3)
rich_para([('我把「多维平衡的活动策划」做成了「线性单向的流水线」。', 11.5, RED, True),
           ('策划的本质是在「商业 / 品牌 / 内容 / 关系 / 体验」多个维度上同时平衡；而五步流水线（拆解→重建→发散→对抗→成型）是线性的，线性结构天然会漏维度。问题 1～4 全是这个错配的不同切面。', 10.5, DARK, False)])

make_table(
    ['问题', '本质', '是否成立', '优先级', '一句话判断'],
    [
        ['1 理论支撑缺失', '推理层缺失：只有机制名词，没有推理链', '成立', 'P0', '清单不是推理，拍脑袋选机制'],
        ['2 指标闭环缺失', '指标是并列KPI，不是传导网络', '成立', 'P0', '只定结果指标，方案退化成冲转化'],
        ['3 角度缺失', '只拆「动机」一个轴，维度不全', '成立', 'P0（最致命）', '单轴拆解必然漏角度'],
        ['4 延展拆分不足', '缺「功能维度完整性」检查', '成立', 'P0', '与问题3同源，奥莱是典型症状'],
        ['5 缺 docx 输出', '交付格式缺陷，非方法论缺陷', '成立', 'P1', '没站在「用户拿到方案后要干嘛」设计'],
    ],
    widths=[3.2, 4.2, 1.6, 2.2, 5.0]
)

# ============ 判断标准 ============
h1('一、判断标准：先定义「活动策划 skill 的第一性」')

para('判断一个活动策划 skill 该不该改，得先回答一个更底层的题：这个 skill 的「第一性」到底是什么。', size=10.5, bold=True)
para('答案是——把用户脑子里「模糊但完整」的意图，无损地翻译成一份「对、全、深、通、用」的方案：', size=10.5)

make_table(
    ['质量维度', '含义', '现有 skill 是否覆盖', '对应问题'],
    [
        ['对（方向对）', '动机和成功判据没搞错', '部分覆盖（Step1-2）', '—'],
        ['全（角度无漏）', '用户要的每个角度都被覆盖，没有 AI 脑补漏项', '缺失', '问题 3、4'],
        ['深（推理有据）', '选心理机制、定策略有学科理论支撑，不是拍脑袋', '缺失', '问题 1'],
        ['通（指标闭环）', '指标之间有传导链路，不是并列 KPI 各自为战', '缺失', '问题 2'],
        ['用（交付可用）', '方案能直接交到老板/客户手里落地（docx）', '缺失', '问题 5'],
    ],
    widths=[2.6, 5.2, 4.4, 2.6]
)

para('结论：策元现有的五步流水线，把「对」这一维做到了一半（只挖了动机，没挖全维度），其余「全 / 深 / 通 / 用」四维基本是空的。用户测出的五个问题，恰好是这四维缺口的外部表现——不是偶然，是结构决定的。', size=10.5, color=RED, bold=True)

# ============ 问题1 ============
h1('二、问题 1：缺心理学 / 社会学 / 营销学理论支撑 —— 成立')

h2('2.1 现状核查')
para('skill 里确实有一份「心理机制清单」（references/creative-inputs.md 第五节），列了 7 个机制：好奇缺口、社交货币、稀缺、损失厌恶、身份认同、从众、情绪峰值，每个配「一句话 + 怎么用 + 案例」。', size=10.5)
para('但这不是「理论支撑」，是「效果名词表」。两者的差别是决定性的。', size=10.5, bold=True)

h2('2.2 第一性原理推理')
para('策划的真正难点，从来不是「知道有哪些心理开关」，而是「针对这个受众、这个动机、这个场景，该选哪个机制、为什么选它、什么条件下它会失效」。', size=10.5)
rich_para([('现有清单只能回答「有哪些」，回答不了「为什么」和「何时失效」。', 10.5, RED, True),
           ('于是创意支点的选择是「拍脑袋挑一个」，不是「推理出来的」。这不是「少了几个学科名词」的问题，是「推理引擎缺失」——清单是静态库存，推理链才是生产力。', 10.5, DARK, False)])

h2('2.3 具体缺什么学科')
make_table(
    ['学科', '该提供的推理依据', '现状'],
    [
        ['消费者心理学', '决策双系统、购买决策链（AIDA/AISAS）、动机理论', '缺失'],
        ['社会心理学', '社会认同理论、从众的边界条件、社会比较、群体极化', '缺失'],
        ['营销学', '定位理论、品牌资产、4C/4P、价格锚点', '缺失'],
        ['社会学', '圈层与亚文化、符号消费、布尔迪厄的区隔', '缺失'],
    ],
    widths=[3.0, 8.0, 4.0]
)
para('举例：清单里的「损失厌恶」只写了「怕失去比想得到更强烈」。但它没说——损失厌恶在「已经拥有后再剥夺」时最强，在「本来就想要但还没得到」时很弱。招商会逼单用「名额流失」有效，发布会造期待用「损失厌恶」就基本失效。缺了这层边界，机制就被乱用。', size=10.5, color=GRAY)

h2('2.4 修改方案')
bullet('把「心理机制清单」升级为「心理机制推理卡」：每个机制挂 4 个字段——理论出处 / 触发条件 / 失效边界 / 反例。', bold=True)
bullet('新增一条「动机→情绪→行为」推理链：先定位受众情绪底层（Step 3 已有），再用推理卡从情绪推到行为，而不是从名词表里挑。')
bullet('仍做成可路由的 reference，遵守「读到哪段停哪段」，不把教科书整本塞进上下文。')

# ============ 问题2 ============
h1('三、问题 2：指标关联延展 + 闭环图 —— 成立')

h2('3.1 现状核查')
para('Step 2 要求「3 个可量化成功判据 + 1 个主判据」，Step 5 有「效果评估」做「目标 vs 实际」复盘。表面看有闭环，但这是「单线闭环」——目标→判据→评估，一条线走到底。', size=10.5)

h2('3.2 第一性原理推理')
para('活动效果是「传导」出来的，不是「并列」出来的。真实链路是这样的：', size=10.5, bold=True)
rich_para([('投入（预算/内容/体验）→ 中间指标（到场率/体验峰值/传播量）→ 结果指标（转化/签约/GMV）→ 长线资产（品牌心智/关系/数据资产）。', 11, BLUE, True)])
para('现有 skill 把「到场率、转化数、传播量」当三个并列 KPI 摆在那里，没讲清它们之间的因果传导。用户说的「一条路走到黑」，根子就在这：', size=10.5)
rich_para([('当主判据只锚定「转化」这一个结果指标、又不建中间指标的传导链时，方案会自然退化——所有资源往转化上堆，品牌演绎、体验创新这些「间接但长线」的环节被砍掉。奥莱商场的病，本质就是这个。', 10.5, RED, True)])

h2('3.3 修改方案')
bullet('新增「指标传导模型」：强制把成功判据拆成一条链——每个结果指标，必须说清由哪几个中间指标驱动；每个中间指标，必须说清由哪项投入驱动。', bold=True)
bullet('把「长线资产指标」（品牌心智、关系、数据资产）纳入判据体系，防止方案短视。')
bullet('产出物层面：强制产出一张「闭环图」——投入→过程→结果→资产，四个节点 + 传导箭头。这张图成为交付物的固定组成部分，且要「根据用户需求」动态生成（不同动机，链路的重心不同）。')

# ============ 问题3 ============
h1('四、问题 3：角度缺失 —— 成立（最致命）')

h2('4.1 现状核查')
para('Step 1 的拆解只有一个轴——「动机」（5 Whys → 钱/权/关系/怕）。字段表列的是「目标、预算、城市、日期、时长、人数、受众」这些执行参数，没有「维度」。', size=10.5)

h2('4.2 第一性原理推理')
para('一场活动需要拆解的，远不止「动机」一个轴。至少还有两个轴被漏掉了：', size=10.5, bold=True)
make_table(
    ['拆解轴', '该拆出什么', '现状'],
    [
        ['目标维度', '商业目标（卖货/拉新）vs 品牌目标（心智/形象/演绎）vs 关系目标（政府/渠道/媒体）', '缺失'],
        ['利益相关方维度', '主办方要什么、受众要什么、参与方（品牌/商家/媒体/KOL）要什么、监管要什么——多方诉求可能互相冲突', '缺失'],
        ['功能维度', '一场活动可同时承担促销 + 品牌演绎 + 内容生产 + 关系经营 + 体验创新 + 数据资产', '缺失'],
    ],
    widths=[3.0, 9.0, 3.0]
)
rich_para([('单轴拆解的后果是结构性的：Step 3 发散由「横纵双轴」驱动，但横纵双轴是在「同一个主轴」上纵深，不会帮你横向换轴。', 10.5, RED, True),
           ('动机只拆出「卖货」一个点，发散就会围着「卖货」打转，把「品牌演绎」「内容破圈」这些用户真正要的角度挤掉。方案和用户需求的错位，从这里开始。', 10.5, DARK, False)])

h2('4.3 修改方案')
bullet('Step 1 增加「需求维度完整性枚举」：把上面三个轴（目标/利益相关方/功能）做成检查清单，逐项过一遍，让用户确认「这次要覆盖哪几个维度」。', bold=True)
bullet('这是对铁律 1「拆到底」的真正兑现——目前只拆到了「动机」，没拆到「维度」。')

# ============ 问题4 ============
h1('五、问题 4：需求延展性拆分不足 —— 成立（与问题 3 同源）')

h2('5.1 奥莱案例的本质')
para('用户举的例子极其精准。奥莱商场做活动，方案全聚焦到「促销折扣转化」上——这是因为「商业转化」是它最显眼、最好量化的功能维度，于是它成了唯一维度。', size=10.5)
rich_para([('「整体演绎」= 品牌演绎 + 体验创新 + 内容传播这三个维度，被「商业转化」一个维度挤掉了。', 10.5, RED, True),
           ('这不是奥莱的错，是 skill 的错：skill 没有「功能维度完整性检查」这个强制环节，所以一旦识别出一个显眼维度，就顺着它一条路走到黑。', 10.5, DARK, False)])

h2('5.2 修改方案')
bullet('在 Step 3 发散前，强制过一遍「活动功能维度清单」，列出这场活动「可能承担的全部功能」，让用户勾选/确认「这次要覆盖哪几个」，未勾选的不进方案，勾选的必须有对应创意承载。', bold=True)
bullet('问题 3 和问题 4 合并处理：核心是给 Step 1/2 加一个「维度完整性」的强制关卡，确保「角度完整」在发散之前就被确认，而不是发散之后才发现漏了。')

# ============ 问题5 ============
h1('六、问题 5：缺 docx 输出 —— 成立（但性质不同）')

h2('6.1 判断')
para('前四个问题是「方法论对不对、全不全」，第五个问题是「方案怎么交到用户手里能直接用」。性质完全不同，但同样是真需求。', size=10.5)
rich_para([('第一性原理：策划方案的终点不是「AI 生成完」，是「交到老板/客户手里、落地执行」。', 10.5, BLUE, True),
           ('中文职场里 docx 是最通用的交付格式。现有 skill 默认输出 markdown 方案，等于把「最后一公里」丢给用户自己转换——这是「自嗨式设计」，只顾自己好写，没顾用户好用。', 10.5, DARK, False)])

h2('6.2 现状与修改方案')
para('现有 skill 的 Step 5「能力检索表」（references/skill-routing.md）只列了「文案写作风格」「文案 AI 味检测」两行，没有「docx 生成」这一行。', size=10.5)
bullet('修改方案：Step 5 能力检索表新增一行「docx 文档生成」，让策元在交付时按能力检索本地 docx 生成能力（如 tencent-docx / html-to-docx / python-docx），借力生成 docx；检索不到就用内置 markdown 兜底并明确告知。', bold=True)

# ============ 统一根因 ============
h1('七、统一根因：线性流水线 vs 多维平衡')

para('把问题 1～4 串起来看，它们不是四个独立 bug，是同一个结构性错配的四个切面：', size=10.5, bold=True)
rich_para([('五步流水线是「线性」的（拆解→重建→发散→对抗→成型），而活动策划是「多维平衡」的（商业/品牌/内容/关系/体验同时平衡）。', 11, RED, True),
           ('线性结构只擅长「往深里走一条线」，不擅长「横着铺满所有维度」。所以：', 10.5, DARK, False)])
bullet('只拆了一个轴（动机）→ 问题 3/4 的角度缺失')
bullet('指标只做并列不建传导 → 问题 2 的一条路走到黑')
bullet('心理机制只列名词不建推理 → 问题 1 的拍脑袋选机制')

para('这也解释了为什么「第一性原理」这条铁律没兜住——铁律 1 说「拆到底」，但实际只拆到了「动机」就停了，没拆到「功能维度 / 利益相关方 / 目标维度」。承诺的深度 > 执行的深度，落差就是这四个问题。', size=10.5, color=GRAY)

# ============ 自身不足 ============
h1('八、我的不足（自我审查）')

para('这五个问题，责任不在 skill 被「用错」，在设计者（我）没想透。逐条认：', size=10.5, bold=True)

bullet('不足一（最核心）：我用了「第一性原理」这个词，却没把它贯彻到「维度拆解」上。', bold=True, color=RED)
para('铁律 1 写「拆到底，只信不可再拆的事实」，但我实际上只拆了「动机」一层，没拆「功能维度」「利益相关方」「目标维度」。这是逻辑不一致——我承诺的比我做的多。问题 3/4 就是这个落差的外显。', size=10, color=DARK, indent=0.6)

bullet('不足二：我把「心理机制」做成了「清单」而不是「推理链」。', bold=True, color=RED)
para('列 7 个名词 + 案例，比建一个「机制→理论→触发条件→失效边界」的推理模型省事得多。但省事的代价是——列名词不等于懂心理学，更不等于能用心理学做推理。这是偷懒。', size=10, color=DARK, indent=0.6)

bullet('不足三：我把「成功判据」做成了「并列 KPI」而不是「传导网络」。', bold=True, color=RED)
para('这暴露我对「活动效果是怎么产生的」理解不深——指标不是并列的三个数字，是投入→过程→结果→资产的传导链。我只给了三个并列数字，等于把链条切成了三截互不相连的碎片。', size=10, color=DARK, indent=0.6)

bullet('不足四：我没站在「用户拿到方案后要干嘛」的立场设计交付。', bold=True, color=RED)
para('策划方案最终要交老板、交客户、要落地，docx 是最通用的格式，我却默认 markdown 就完事。这是「自嗨式设计」——只顾自己产出方便，没顾用户使用方便。问题 5 的根子在这。', size=10, color=DARK, indent=0.6)

bullet('不足五：我把「活动策划」过度收敛成了一个「线性流程」。', bold=True, color=RED)
para('这是五个问题里 1～4 的共同根因。真实策划是在多个维度上「同时平衡」，我却把它建模成「一条线走到底」。线性模型天然漏维度，这是建模方式错了，不是执行不到位。', size=10, color=DARK, indent=0.6)

# ============ 修改优先级 ============
h1('九、修改优先级与执行建议')

make_table(
    ['优先级', '问题', '改什么', '为什么这个顺序'],
    [
        ['P0-1', '3 + 4 角度/维度', 'Step 1/2 加「维度完整性」强制关卡', '维度错漏最致命，方案会整体跑偏，必须先堵'],
        ['P0-2', '2 指标闭环', '加「指标传导模型」+ 闭环图产出', '指标不通，方案退化成冲转化，奥莱病'],
        ['P0-3', '1 理论支撑', '清单升级为「推理卡」', '推理有据，创意才不是拍脑袋'],
        ['P1', '5 docx 输出', 'Step 5 能力表加 docx 生成', '交付格式，外围但必加'],
    ],
    widths=[1.6, 2.4, 6.6, 4.4]
)

para('执行方式：按策元自己的铁律「你不点头，不往下走」，逐关确认，不擅自改。', size=10.5)

# ============ 收尾 ============
h1('十、收尾判词')

rich_para([('这五个问题不是「skill 没做好」，是「我把它想浅了」。', 11.5, RED, True)])
para('真正要改的不是加几个文件、补几段话，是换一个建模前提：从「线性流水线」换成「多维平衡」。前提不换，补再多 patch 都是给错误的骨架打补丁。', size=11)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
run = p.add_run('—— 牛牛 · 2026-08-14')
set_font(run, size=10, color=GRAY)

doc.save(r'E:\我的AI项目聚集地\策划skill\ceyuan五问题诊断报告.docx')
print('OK saved')
